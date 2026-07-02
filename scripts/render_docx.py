# -*- coding: utf-8 -*-
"""
render_docx.py — 把 pandoc 生成的 docx 后处理成符合中文学术规范的格式。

用法：
  python render_docx.py input.docx output.docx [--line-spacing 1.5]

负责：
  1. 清理 pandoc frontmatter 残留（Author/Date/重复标题）
  2. 标题黑体居中、小标题黑体、正文宋体小四首行缩进
  3. 全篇黑色（去掉 pandoc 标题默认蓝色）
  4. 脚注正文：宋体小五 + Times New Roman 9pt
  5. 文档默认字体：宋体 + Times New Roman（防止等线混入）
  6. 参考文献：宋体五号、序号左顶格
  7. 参考文献前分页（单起一页）
  8. 脚注编号：每页重新编号 + 圈码 ①②③（decimalEnclosedCircle）
  9. 页边距 A4

依赖：python-docx
"""
import argparse
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_run_font(run, latin, east, size, bold=False):
    """设置 run 字体并强制黑色。"""
    run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), east)
    rFonts.set(qn('w:ascii'), latin)
    rFonts.set(qn('w:hAnsi'), latin)
    run.font.size = Pt(size)
    run.bold = bold
    # 强制黑色，覆盖 pandoc 标题样式里的蓝色
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_doc_defaults(doc):
    """设置文档默认字体为宋体 + Times New Roman，防止等线混入。"""
    styles_el = doc.styles.element
    docDefaults = styles_el.find(qn('w:docDefaults'))
    if docDefaults is None:
        docDefaults = OxmlElement('w:docDefaults')
        styles_el.insert(0, docDefaults)
    rPrDefault = docDefaults.find(qn('w:rPrDefault'))
    if rPrDefault is None:
        rPrDefault = OxmlElement('w:rPrDefault')
        docDefaults.append(rPrDefault)
    rPr = rPrDefault.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        rPrDefault.append(rPr)
    # 设置默认字体
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:cs'), 'Times New Roman')


def set_style_rpr(doc, style_id, east, latin, size_pt):
    """给指定样式注入 rPr（字体+字号），用于脚注等非正文段落。"""
    styles_el = doc.styles.element
    for style in styles_el.findall(qn('w:style')):
        if style.get(qn('w:styleId')) == style_id:
            rPr = style.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                # rPr 应在 pPr 之后
                pPr = style.find(qn('w:pPr'))
                if pPr is not None:
                    pPr.addnext(rPr)
                else:
                    style.append(rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:ascii'), latin)
            rFonts.set(qn('w:hAnsi'), latin)
            rFonts.set(qn('w:eastAsia'), east)
            rFonts.set(qn('w:cs'), latin)
            sz = rPr.find(qn('w:sz'))
            if sz is None:
                sz = OxmlElement('w:sz')
                rPr.append(sz)
            sz.set(qn('w:val'), str(int(size_pt * 2)))
            szCs = rPr.find(qn('w:szCs'))
            if szCs is None:
                szCs = OxmlElement('w:szCs')
                rPr.append(szCs)
            szCs.set(qn('w:val'), str(int(size_pt * 2)))
            return True
    return False


def set_footnote_per_page_restart(doc):
    """设置脚注：每页重新编号 + 圈码 ①②③。

    numFmt='decimalEnclosedCircle' 是 ECMA-376 标准合法值，对应 ①②③（1-20）。
    注：部分 Word 版本可能不渲染圈码或忽略 eachPage 重编号，
    这是已知的兼容性问题，需结合其他手段进一步解决。
    """
    settings = doc.settings.element
    for old in settings.findall(qn('w:footnotePr')):
        settings.remove(old)
    fnPr = OxmlElement('w:footnotePr')
    numRestart = OxmlElement('w:numRestart')
    numRestart.set(qn('w:val'), 'eachPage')
    fnPr.append(numRestart)
    numStart = OxmlElement('w:numStart')
    numStart.set(qn('w:val'), '1')
    fnPr.append(numStart)
    numFmt = OxmlElement('w:numFmt')
    numFmt.set(qn('w:val'), 'decimalEnclosedCircle')
    fnPr.append(numFmt)
    settings.insert(0, fnPr)


def is_section_heading(text):
    markers = ['一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、']
    return any(text.startswith(m) for m in markers)


def format_document(doc, line_spacing=1.5):
    title_text = None
    for p in doc.paragraphs:
        if p.style and p.style.name == 'Title':
            title_text = p.text.strip()
            break

    to_remove = []
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else ''
        text = p.text.strip()

        # 清理 frontmatter 残留
        if style in ('Author', 'Date') and not text:
            to_remove.append(i)
            continue
        if style == 'Date':
            to_remove.append(i)
            continue
        if style == 'Heading 1' and text == title_text:
            to_remove.append(i)
            continue

        if style == 'Title':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = line_spacing
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.first_line_indent = Pt(0)
            for r in p.runs:
                set_run_font(r, 'Times New Roman', '黑体', 18, bold=True)

        elif style.startswith('Heading 1') or text == '参考文献':
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = line_spacing
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Pt(0)
            # 参考文献单起一页
            if text == '参考文献':
                p.paragraph_format.page_break_before = True
            for r in p.runs:
                set_run_font(r, 'Times New Roman', '黑体', 14, bold=True)

        elif style.startswith('Heading 2') or is_section_heading(text):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = line_spacing
            p.paragraph_format.first_line_indent = Pt(24)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            for r in p.runs:
                set_run_font(r, 'Times New Roman', '黑体', 14, bold=True)

        elif text.startswith('[') and ']' in text[:6]:
            # 参考文献条目：宋体五号
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                set_run_font(r, 'Times New Roman', '宋体', 10.5)

        else:
            # 正文：宋体小四
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = line_spacing
            p.paragraph_format.first_line_indent = Pt(24)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                set_run_font(r, 'Times New Roman', '宋体', 12)

    for i in reversed(to_remove):
        el = doc.paragraphs[i]._element
        el.getparent().remove(el)

    # 页边距 A4
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.18)
    sec.right_margin = Cm(3.18)

    # 文档默认字体（防等线）
    set_doc_defaults(doc)
    # 脚注正文样式：宋体小五 + TNR 9pt
    set_style_rpr(doc, 'FootnoteText', '宋体', 'Times New Roman', 9)
    set_style_rpr(doc, 'FootnoteBlockText', '宋体', 'Times New Roman', 9)
    # 脚注编号：每页重编号 + 圈码
    set_footnote_per_page_restart(doc)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('input')
    ap.add_argument('output')
    ap.add_argument('--line-spacing', type=float, default=1.5)
    args = ap.parse_args()

    doc = Document(args.input)
    format_document(doc, args.line_spacing)
    doc.save(args.output)
    print(f'已格式化: {args.output}')


if __name__ == '__main__':
    main()
